import os
import nltk
from nltk.tokenize import word_tokenize
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
import markdown
from openpyxl import load_workbook
from multiprocessing import Process, Manager

nltk_data_path = "./scripts/nltk_data"

if os.path.exists(nltk_data_path):
    nltk.data.path.append(nltk_data_path)


class DocumentHandler():
    @staticmethod
    def tokenize_text(text):
        tokens = word_tokenize(text)
        return tokens

    @staticmethod
    def tokenize_html(html):
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        tokens = word_tokenize(text)
        return tokens

    @staticmethod
    def tokenize_pdf(pdf_path):
        tokens = []
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                tokens.extend(word_tokenize(text))
        return tokens

    @staticmethod
    def extract_paragraph_from_docx(paragraph):
        return word_tokenize(paragraph.text)+['\n']

    @staticmethod
    def extract_table_from_docx(table):
        tokens = []
        tokens.extend('\n\n')
        header_printed = False

        for row in table.rows:
            row_tokens = []
            if not header_printed:
                header_line = ['-' * (len(str(cell.text).strip()) + 2) for cell in row.cells]
                header_printed = True

            for cell in row.cells:
                row_tokens.extend(word_tokenize(str(cell.text).strip()))
                row_tokens.append(' | ')

            tokens.extend(row_tokens)
            tokens.append('\n')

        if header_printed:
            tokens.extend(['-'.join(header_line), '\n'])
        tokens.extend('\n\n')
        return tokens

    @staticmethod
    def create_docx_element_map(doc):
        element_map = {}
        for paragraph in doc.paragraphs:
            element_map[paragraph._element] = paragraph
        for table in doc.tables:
            element_map[table._element] = table
        return element_map

    @staticmethod
    def get_sub_tokens(pid, element_map, sub_el_list, q):
        for element in sub_el_list:
            sub_tokens = []
            if element in element_map:
                obj = element_map[element]
                if obj._element.tag.endswith('p'):
                    sub_tokens = DocumentHandler.extract_paragraph_from_docx(obj)
                elif obj._element.tag.endswith('tbl'):
                    sub_tokens = DocumentHandler.extract_table_from_docx(obj)
            q.put((pid, sub_tokens))

        q.put(None)

    @staticmethod
    def tokenize_docx(docx_path, num_cores=8):
        doc = Document(docx_path)
        tmp_tokens = []

        el_list = []
        for element in doc.element.body:
            el_list.append(element)
        if os.cpu_count() is None:
            return []
        num_cores = min(num_cores, os.cpu_count()//2)
        num_cores = min(8, min(num_cores, len(el_list)))
        num_cores = max(num_cores, 1)
        chunk_sz = len(el_list)//num_cores
        for i in range(num_cores):
            tmp_tokens.append([])
        processes = []
        q = Manager().Queue()
        element_map = DocumentHandler.create_docx_element_map(doc)
        for i in range(num_cores):
            st = i*chunk_sz
            en = (i+1)*chunk_sz
            if i == num_cores-1:
                en = len(el_list)
            p = Process(target=DocumentHandler.get_sub_tokens, args=(i, element_map, el_list[st:en], q))
            processes.append(p)
            p.start()
        tokens = []
        task_finish_cnt = 0
        tmp_cnt = 0
        while task_finish_cnt < num_cores:
            tmp = q.get()
            tmp_cnt += 1
            if tmp is None:
                task_finish_cnt += 1
                if task_finish_cnt == num_cores:
                    break
                continue
            else:
                id, sub_tokens = tmp
            if tmp_cnt % 1000 == 0:
                print(tmp_cnt)
            tmp_tokens[id].extend(sub_tokens)
        for sub_tokens in tmp_tokens:
            tokens.extend(sub_tokens)
        return tokens

    @staticmethod
    def tokenize_md(md_path):
        with open(md_path, 'r', encoding='utf-8', errors="ignore") as file:
            md_text = file.read()
            html_text = markdown.markdown(md_text)
            tokens = DocumentHandler.tokenize_html(html_text)
        return tokens

    @staticmethod
    def tokenize_xlsx(file_path):
        workbook = load_workbook(filename=file_path)
        table_list = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            merged_ranges = sheet.merged_cells.ranges
            max_column = sheet.max_column
            merged_cells_map = {}
            for merged_range in merged_ranges:
                min_col, min_row, max_col, max_row = merged_range.bounds
                for row in range(min_row, max_row + 1):
                    for col in range(min_col, max_col + 1):
                        merged_cells_map[(row, col)] = (min_row, min_col)

            headers = [cell.value for cell in next(sheet.iter_rows(max_col=max_column))]
            row_list = ["| " + " | ".join(headers) + " |\n"]
            row_list.append("| " + " | ".join(['---'] * len(headers)) + " |\n")

            for row_idx, row in enumerate(sheet.iter_rows(min_row=2, max_col=max_column, values_only=True), start=2):
                row_data = []
                for idx, cell_value in enumerate(row):
                    cell_position = (row_idx, idx + 1)
                    if cell_position in merged_cells_map:
                        top_left_cell_pos = merged_cells_map[cell_position]
                        cell_value = sheet.cell(row=top_left_cell_pos[0], column=top_left_cell_pos[1]).value
                    row_data.append(str(cell_value) if cell_value is not None else '')
                row_list.append("| " + " | ".join(row_data) + " |\n")
            table_list.append(row_list)
        return table_list

    @staticmethod
    def split_into_paragraphs(words, max_paragraph_length=1024):
        if max_paragraph_length == -1:
            return words
        paragraphs = []
        current_paragraph = ""
        for word in words:
            if max_paragraph_length is  None or len(current_paragraph) + len(word) <= max_paragraph_length:
                current_paragraph += word
            else:
                paragraphs.append(current_paragraph)
                current_paragraph = word

        if current_paragraph:
            paragraphs.append(current_paragraph.strip())
        return paragraphs

    def get_content_list_from_file(file_path, max_paragraph_length=1024, num_cores=8):
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors="ignore") as file:
                text = file.read()
                tokens = DocumentHandler.tokenize_text(text)
                paragraphs = DocumentHandler.split_into_paragraphs(tokens, max_paragraph_length)
                return paragraphs
        elif file_extension == '.html':
            with open(file_path, 'r', encoding='utf-8', errors="ignore") as file:
                html_text = file.read()
                tokens = DocumentHandler.tokenize_html(html_text)
                paragraphs = DocumentHandler.split_into_paragraphs(tokens, max_paragraph_length)
                return paragraphs
        elif file_extension == '.pdf':
            tokens = DocumentHandler.tokenize_pdf(file_path)
            paragraphs = DocumentHandler.split_into_paragraphs(tokens, max_paragraph_length)
            return paragraphs
        elif file_extension == '.docx':
            tokens = DocumentHandler.tokenize_docx(file_path, num_cores)
            paragraphs = DocumentHandler.split_into_paragraphs(tokens, max_paragraph_length)
            return paragraphs
        elif file_extension == '.md':
            tokens = DocumentHandler.tokenize_md(file_path)
            paragraphs = DocumentHandler.split_into_paragraphs(tokens, max_paragraph_length)
            return paragraphs
        elif file_extension == '.xlsx':
            table_list = DocumentHandler.tokenize_xlsx(file_path)
            paragraphs = []
            for table in table_list:
                paragraphs += DocumentHandler.split_into_paragraphs(table, max_paragraph_length)
            return paragraphs
        else:
            return []
